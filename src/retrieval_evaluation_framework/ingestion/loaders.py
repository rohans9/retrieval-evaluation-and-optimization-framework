"""Document ingestion implementations."""

from __future__ import annotations

from csv import reader
from importlib import metadata
import hashlib
import re
import unicodedata
from pathlib import Path
from typing import Any

from docx import Document as DocxDocument, text
from pypdf import PdfReader

from retrieval_evaluation_framework.ingestion.base import DocumentLoader
from retrieval_evaluation_framework.logging import get_logger
from retrieval_evaluation_framework.models import Document, FileType

LOGGER = get_logger(component="ingestion")

LIGATURE_MAP = {
    "\ufb00": "ff",
    "\ufb01": "fi",
    "\ufb02": "fl",
    "\ufb03": "ffi",
    "\ufb04": "ffl",
    "\ufb05": "st",
    "\ufb06": "st",
}

def _build_document_id(file_path: Path, text: str) -> str:
    # Use filename + content so IDs stay stable across machines and repo locations.
    digest = hashlib.sha256(f"{file_path.name}::{text}".encode()).hexdigest()
    return digest[:16]


def _base_metadata(file_path: Path) -> dict[str, Any]:
    stat = file_path.stat()
    return {
        "filename": file_path.name,
        "extension": file_path.suffix.lower(),
        "size_bytes": stat.st_size,
        "modified_at": stat.st_mtime,
    }

def normalize_pdf_text(text: str) -> str:
    """Normalize extracted PDF text."""

    # Unicode normalization
    text = unicodedata.normalize("NFKC", text)

    # Replace common ligatures
    for ligature, replacement in LIGATURE_MAP.items():
        text = text.replace(ligature, replacement)

    # Remove replacement characters
    text = text.replace("\ufffd", "")

    # Remove soft hyphens
    text = text.replace("\u00ad", "")

    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")

    # Collapse spaces
    text = re.sub(r"[ \t]+", " ", text)

    # Collapse excessive blank lines
    text = re.sub(r"\n{3,}", "\n\n", text)

    # Remove invisible Unicode characters
    text = re.sub(r"[\u200b-\u200f\u202a-\u202e\u2060\ufeff]", "", text)

    # Remove duplicate spaces around newlines
    text = re.sub(r" *\n *", "\n", text)

    return text.strip()

class TextLoader(DocumentLoader):
    """Loader for plain-text documents."""

    supported_extensions = (".txt",)

    def load(self, file_path: Path) -> Document:
        text = file_path.read_text(encoding="utf-8")
        return Document(
            id=_build_document_id(file_path, text),
            title=file_path.stem,
            text=text,
            source=str(file_path),
            file_type=FileType.TXT,
            metadata=_base_metadata(file_path),
        )


class MarkdownLoader(DocumentLoader):
    """Loader for Markdown documents."""

    supported_extensions = (".md",)

    def load(self, file_path: Path) -> Document:
        text = file_path.read_text(encoding="utf-8")
        return Document(
            id=_build_document_id(file_path, text),
            title=file_path.stem,
            text=text,
            source=str(file_path),
            file_type=FileType.MARKDOWN,
            metadata=_base_metadata(file_path),
        )


class PdfLoader(DocumentLoader):
    """Loader for PDF documents."""

    supported_extensions = (".pdf",)

    def load(self, file_path: Path) -> Document:
        reader = PdfReader(str(file_path))
        page_text = []

        for page_number, page in enumerate(reader.pages, start=1):
            try:
                extracted = page.extract_text(extraction_mode="layout") or ""
            except TypeError:
                # Older versions of pypdf don't support extraction_mode
                extracted = page.extract_text() or ""

            extracted = normalize_pdf_text(extracted)

            if len(extracted.strip()) < 20:
                LOGGER.warning(
                    "poor_pdf_extraction",
                    path=str(file_path),
                    page=page_number,
                )

            page_text.append(extracted)

        text = "\f".join(page_text)

        metadata = _base_metadata(file_path)

        metadata["page_count"] = len(reader.pages)
        metadata["extraction"] = {
            "engine": "pypdf",
            "pages": len(reader.pages),
            "characters": len(text),
        }
        if reader.metadata:
            metadata["pdf_metadata"] = {
                str(key): str(value) for key, value in reader.metadata.items()
            }
        return Document(
            id=_build_document_id(file_path, text),
            title=file_path.stem,
            text=text,
            source=str(file_path),
            file_type=FileType.PDF,
            metadata=metadata,
        )


class DocxLoader(DocumentLoader):
    """Loader for DOCX documents."""

    supported_extensions = (".docx",)

    def load(self, file_path: Path) -> Document:
        docx_document = DocxDocument(str(file_path))
        paragraphs = [
            paragraph.text for paragraph in docx_document.paragraphs if paragraph.text.strip()
        ]
        text = "\n".join(paragraphs)
        metadata = _base_metadata(file_path)
        metadata["paragraph_count"] = len(paragraphs)
        return Document(
            id=_build_document_id(file_path, text),
            title=file_path.stem,
            text=text,
            source=str(file_path),
            file_type=FileType.DOCX,
            metadata=metadata,
        )


class DocumentIngestor:
    """Coordinator for document discovery and loader dispatch."""

    def __init__(self, loaders: list[DocumentLoader] | None = None) -> None:
        """Initialize the ingestor.

        Args:
            loaders: Optional list of loader instances.
        """
        self.loaders = loaders or [PdfLoader(), DocxLoader(), TextLoader(), MarkdownLoader()]

    def discover_files(self, path: Path, recursive: bool = True) -> list[Path]:
        """Discover supported files under a path.

        Args:
            path: Source file or directory.
            recursive: Whether to recurse into subdirectories.

        Returns:
            Sorted list of supported file paths.
        """
        if path.is_file():
            return [path] if self._get_loader(path) else []

        pattern = "**/*" if recursive else "*"
        discovered = [
            candidate
            for candidate in path.glob(pattern)
            if candidate.is_file() and self._get_loader(candidate)
        ]
        for candidate in discovered:
            LOGGER.info("discovered_file", path=str(candidate))
        return sorted(discovered)

    def ingest_file(self, file_path: Path) -> Document:
        """Parse a single file into a document.

        Args:
            file_path: File to parse.

        Returns:
            Parsed document.
        """
        loader = self._get_loader(file_path)
        if loader is None:
            msg = f"Unsupported file type: {file_path.suffix}"
            raise ValueError(msg)
        document = loader.load(file_path)
        LOGGER.info("parsed_file", path=str(file_path), document_id=document.id)
        return document

    def ingest_directory(self, directory: Path, recursive: bool = True) -> list[Document]:
        """Parse all supported files in a directory.

        Args:
            directory: Directory to traverse.
            recursive: Whether to recurse into nested directories.

        Returns:
            Parsed documents.
        """
        documents: list[Document] = []
        for file_path in self.discover_files(directory, recursive=recursive):
            try:
                documents.append(self.ingest_file(file_path))
            except Exception as error:
                LOGGER.warning("failed_to_parse_file", path=str(file_path), error=str(error))
        return documents

    def _get_loader(self, file_path: Path) -> DocumentLoader | None:
        for loader in self.loaders:
            if loader.supports(file_path):
                return loader
        return None
